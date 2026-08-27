import os
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def generate_diagram(image_path):
    fig, ax = plt.subplots(figsize=(8.5, 3.8), dpi=300)
    ax.axis('off')

    # Colors
    bg_color = "#f8fafc"
    fig.patch.set_facecolor(bg_color)
    
    c_user = "#1e293b"      # Slate 800
    c_fe = "#0f766e"        # Teal 700
    c_be = "#0369a1"        # Sky 700
    c_ai = "#165823"        # Brand Green
    c_out = "#7c3aed"       # Purple 600

    def draw_box(ax, x, y, w, h, title, subtitle, color, text_color="white"):
        rect = patches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.04,rounding_size=0.08",
                                      facecolor=color, edgecolor="#cbd5e1", linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h*0.62, title, color=text_color, fontsize=10, fontweight='bold',
                ha='center', va='center', fontfamily='sans-serif')
        ax.text(x + w/2, y + h*0.28, subtitle, color=text_color, fontsize=7.5,
                ha='center', va='center', fontfamily='sans-serif', alpha=0.9)

    # 1. User & Client Layer
    draw_box(ax, 0.02, 0.55, 0.20, 0.35, "E-Commerce Founder", "Store Catalog & Goals", c_user)
    draw_box(ax, 0.02, 0.10, 0.20, 0.35, "React 18 + TS Frontend", "Vite, Tailwind, Context", c_fe)

    # 2. Backend Orchestration Layer
    draw_box(ax, 0.29, 0.55, 0.26, 0.35, "FastAPI Backend Core", "Multi-Tenant API Gateway", c_be)
    draw_box(ax, 0.29, 0.10, 0.26, 0.35, "Supabase PostgreSQL", "Row-Level Security (RLS)", "#334155")

    # 3. AI Agent Intelligence Engine
    draw_box(ax, 0.62, 0.55, 0.36, 0.35, "Google Gemini 3.6 Flash", "Margin-Aware Context Ingestion", c_ai)
    draw_box(ax, 0.62, 0.10, 0.17, 0.35, "Trend Radar", "Live Trends RSS", "#d97706")
    draw_box(ax, 0.81, 0.10, 0.17, 0.35, "Brand Guardrails", "Policy Validator", "#be123c")

    # Connectors / Arrows
    def draw_arrow(ax, x1, y1, x2, y2, label=""):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="->,head_width=0.3,head_length=0.4",
                                    color="#475569", lw=1.8))
        if label:
            ax.text((x1+x2)/2, (y1+y2)/2 + 0.03, label, fontsize=7, color="#475569",
                    ha='center', va='bottom', fontweight='bold')

    draw_arrow(ax, 0.12, 0.55, 0.12, 0.45)
    draw_arrow(ax, 0.22, 0.27, 0.29, 0.27, "Auth & RLS")
    draw_arrow(ax, 0.22, 0.72, 0.29, 0.72, "JSON Payload")
    draw_arrow(ax, 0.55, 0.72, 0.62, 0.72, "Grounded Prompt")
    draw_arrow(ax, 0.70, 0.45, 0.70, 0.55)
    draw_arrow(ax, 0.89, 0.45, 0.89, 0.55)

    plt.tight_layout()
    plt.savefig(image_path, bbox_inches='tight', dpi=300)
    plt.close()

def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = Document()

    # Configure Margins (0.7 inch all sides for clean 3-page layout)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Color Palette Constants
    c_brand_navy = RGBColor(15, 23, 42)     # #0f172a
    c_brand_green = RGBColor(22, 88, 35)    # #165823
    c_text_dark = RGBColor(30, 41, 59)      # #1e293b
    c_text_muted = RGBColor(100, 116, 139)  # #64748b

    # Base Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = c_text_dark

    # Title Banner Table
    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_cell = title_table.cell(0, 0)
    title_cell.width = Inches(7.1)
    set_cell_shading(title_cell, "0f172a") # Navy
    set_cell_margins(title_cell, top=180, bottom=180, left=240, right=240)

    p_tag = title_cell.paragraphs[0]
    p_tag.paragraph_format.space_before = Pt(0)
    p_tag.paragraph_format.space_after = Pt(2)
    run_tag = p_tag.add_run("INTERNSHIP TECHNICAL PROJECT REPORT · E-COMMERCE AI")
    run_tag.font.size = Pt(8.5)
    run_tag.font.bold = True
    run_tag.font.color.rgb = RGBColor(74, 222, 128) # Light Green

    p_title = title_cell.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("MarketPilot AI: Autonomous Margin-Aware Marketing Copilot")
    run_title.font.size = Pt(17)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(255, 255, 255)

    p_sub = title_cell.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(0)
    run_sub = p_sub.add_run("Architecture, Agentic Working Mechanism, Multi-Channel Strategy & LLM Orchestration")
    run_sub.font.size = Pt(9.5)
    run_sub.font.color.rgb = RGBColor(203, 213, 225)

    # Spacing after banner
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(4)
    p_spacer.paragraph_format.space_after = Pt(4)

    # 1. Executive Summary
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(3)
    r_h1 = h1.add_run("1. Executive Summary & Core Motivation")
    r_h1.font.size = Pt(12)
    r_h1.font.bold = True
    r_h1.font.color.rgb = c_brand_green

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(4)
    p1.paragraph_format.line_spacing = 1.15
    p1.add_run(
        "MarketPilot AI is an autonomous, full-stack marketing operating system built to solve a critical dilemma in direct-to-consumer (D2C) e-commerce: generic AI copywriting tools generate superficial content disconnected from real unit economics, profit margins, and viral consumer search velocity. "
        "Standard LLM prompts frequently promote low-margin items that drain advertising budgets without delivering real profitability. MarketPilot AI bridges this fundamental gap by combining automated margin categorization, live trend intelligence, deterministic brand guardrails, and Google Gemini 3.6 Flash orchestration into a unified daily marketing copilot."
    )

    # 2. Technology Stack & Architecture
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(3)
    r_h2 = h2.add_run("2. Technology Stack & Infrastructure")
    r_h2.font.size = Pt(12)
    r_h2.font.bold = True
    r_h2.font.color.rgb = c_brand_green

    # Tech Stack Table
    tech_table = doc.add_table(rows=5, cols=2)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = [("Layer / Component", "Technologies & Implementation Role"),
               ("Frontend Client", "React 18, TypeScript, Vite, Tailwind CSS, Lucide Icons, SPA routing (Vercel)"),
               ("Backend Core", "FastAPI (Python 3.12+), Uvicorn ASGI server, Pydantic v2 data contracts (Render)"),
               ("Database & Security", "Supabase PostgreSQL with strict Row-Level Security (RLS) & multi-tenant isolation"),
               ("AI Engine & LLM", "Google Gemini 3.6 Flash via google-genai SDK with deterministic guardrail layers")]

    for row_idx, (col1, col2) in enumerate(headers):
        cell1, cell2 = tech_table.cell(row_idx, 0), tech_table.cell(row_idx, 1)
        cell1.width, cell2.width = Inches(2.2), Inches(4.9)
        set_cell_margins(cell1, top=60, bottom=60, left=100, right=100)
        set_cell_margins(cell2, top=60, bottom=60, left=100, right=100)
        
        if row_idx == 0:
            set_cell_shading(cell1, "f1f5f9")
            set_cell_shading(cell2, "f1f5f9")
            p_c1 = cell1.paragraphs[0]
            p_c1.paragraph_format.space_before = Pt(0)
            p_c1.paragraph_format.space_after = Pt(0)
            r = p_c1.add_run(col1)
            r.font.bold = True
            r.font.size = Pt(9)
            
            p_c2 = cell2.paragraphs[0]
            p_c2.paragraph_format.space_before = Pt(0)
            p_c2.paragraph_format.space_after = Pt(0)
            r = p_c2.add_run(col2)
            r.font.bold = True
            r.font.size = Pt(9)
        else:
            set_cell_shading(cell1, "ffffff" if row_idx % 2 == 1 else "f8fafc")
            set_cell_shading(cell2, "ffffff" if row_idx % 2 == 1 else "f8fafc")
            p_c1 = cell1.paragraphs[0]
            p_c1.paragraph_format.space_before = Pt(0)
            p_c1.paragraph_format.space_after = Pt(0)
            r = p_c1.add_run(col1)
            r.font.bold = True
            r.font.size = Pt(8.5)
            
            p_c2 = cell2.paragraphs[0]
            p_c2.paragraph_format.space_before = Pt(0)
            p_c2.paragraph_format.space_after = Pt(0)
            r = p_c2.add_run(col2)
            r.font.size = Pt(8.5)

    # 3. Agentic Working Mechanism & Diagram
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(3)
    r_h3 = h3.add_run("3. End-to-End AI Agent Working Mechanism")
    r_h3.font.size = Pt(12)
    r_h3.font.bold = True
    r_h3.font.color.rgb = c_brand_green

    p_diag_intro = doc.add_paragraph()
    p_diag_intro.paragraph_format.space_after = Pt(4)
    p_diag_intro.paragraph_format.line_spacing = 1.15
    p_diag_intro.add_run(
        "The agent follows a deterministic multi-stage orchestration pipeline. It first extracts catalog profit margins and user budget constraints, enriches them with real-time viral signals, prompts Google Gemini 3.6 Flash with structured JSON schemas, and passes generated copy through a real-time policy safety filter:"
    )

    # Generate and Embed Diagram Image
    img_path = "C:\\Users\\danis\\Desktop\\Internship Project\\architecture_diagram.png"
    generate_diagram(img_path)
    
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(2)
    p_img.paragraph_format.space_after = Pt(2)
    doc.add_picture(img_path, width=Inches(6.9))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(6)
    r_cap = p_cap.add_run("Figure 1: MarketPilot AI Autonomous Agent Architecture & Data Pipeline")
    r_cap.font.size = Pt(8)
    r_cap.font.italic = True
    r_cap.font.color.rgb = c_text_muted

    # 4. Key Agent Innovations
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(3)
    r_h4 = h4.add_run("4. Core AI Agent Capabilities & Innovations")
    r_h4.font.size = Pt(12)
    r_h4.font.bold = True
    r_h4.font.color.rgb = c_brand_green

    innovations = [
        ("Unit Economics & Margin Engine:", "Automatically analyzes product cost vs. retail prices to assign margin tiers (High >65%, Medium 40-65%, Low <40%). The AI prioritizes ad spend and video hooks exclusively on high-margin Hero SKUs to ensure positive Return on Ad Spend (ROAS)."),
        ("Live Trend Intelligence & Virality Radar:", "Continuously ingests trending search queries and viral TikTok/Google Trends hashtags with automated confidence scoring (0-100%) to pair real-time viral search momentum with corresponding catalog items."),
        ("Transparent Grounding & Attribution Panel:", "Every generated output features a prominent 'Based on:' panel detailing the source SKU, profit margin, active discount offer, trend angle, and brand voice traits. It alerts the user with warnings if pain points or features are missing."),
        ("Real-Time Deterministic Brand Guardrails:", "Evaluates copywriting against Brand Kit voice rules and prohibited words (e.g. 'miracle cure', '100% guaranteed'), preventing ad account suspensions and false medical claims in real time."),
        ("Multi-Channel Copywriting Studio:", "Generates platform-native content across 5 distinct channels: TikTok Video Scripts (with 3-second split-screen timestamped cues), Instagram Carousels, Paid Direct-Response Ads, Email Newsletters, and 1-Click WhatsApp VIP Broadcasts."),
        ("Autonomous Daily Copilot ('What to Post Today'):", "Analyzes the calendar date and product stock levels on overview load, recommending the exact high-converting hook and channel strategy to execute immediately.")
    ]

    for title, desc in innovations:
        p_inn = doc.add_paragraph()
        p_inn.paragraph_format.space_before = Pt(0)
        p_inn.paragraph_format.space_after = Pt(3.5)
        p_inn.paragraph_format.line_spacing = 1.15
        r_t = p_inn.add_run(f"• {title} ")
        r_t.font.bold = True
        r_t.font.size = Pt(9.5)
        r_t.font.color.rgb = c_brand_navy
        r_d = p_inn.add_run(desc)
        r_d.font.size = Pt(9.5)

    # 5. Security, Deployment & Verification
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(8)
    h5.paragraph_format.space_after = Pt(3)
    r_h5 = h5.add_run("5. Security, Production Hardening & Global Multi-Currency")
    r_h5.font.size = Pt(12)
    r_h5.font.bold = True
    r_h5.font.color.rgb = c_brand_green

    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(4)
    p5.paragraph_format.line_spacing = 1.15
    p5.add_run(
        "MarketPilot AI is fully hardened for production deployment: (1) Supabase Authentication enforces strict password encryption and secure OTP confirmation without unverified bypasses; (2) Demo authentication tokens are strictly disabled in production environments; (3) The application supports real-time multi-currency conversion across PKR (Rs.), USD ($), AED, SAR, GBP, EUR, CAD, and INR; (4) Frontend is packaged for zero-downtime deployment on Vercel, and the FastAPI service is structured with render.yaml for seamless Render cloud hosting."
    )

    # 6. Conclusion Callout Box
    concl_table = doc.add_table(rows=1, cols=1)
    concl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_cell = concl_table.cell(0, 0)
    c_cell.width = Inches(7.1)
    set_cell_shading(c_cell, "f0fdf4") # Emerald 50
    set_cell_margins(c_cell, top=120, bottom=120, left=160, right=160)

    p_c_title = c_cell.paragraphs[0]
    p_c_title.paragraph_format.space_before = Pt(0)
    p_c_title.paragraph_format.space_after = Pt(2)
    r_ct = p_c_title.add_run("Conclusion & Impact")
    r_ct.font.bold = True
    r_ct.font.size = Pt(10)
    r_ct.font.color.rgb = c_brand_green

    p_c_body = c_cell.add_paragraph()
    p_c_body.paragraph_format.space_before = Pt(0)
    p_c_body.paragraph_format.space_after = Pt(0)
    p_c_body.paragraph_format.line_spacing = 1.15
    r_cb = p_c_body.add_run(
        "This internship project delivers a complete, market-ready AI agent that replaces hours of manual campaign planning with 15-second, unit-economics-grounded marketing execution. By aligning Gemini 3.6 Flash with real profit margins, live breakout trends, and strict brand guardrails, MarketPilot AI establishes a new standard for modern autonomous e-commerce growth."
    )
    r_cb.font.size = Pt(9)
    r_cb.font.color.rgb = RGBColor(22, 101, 52)

    output_path = "C:\\Users\\danis\\Desktop\\Internship Project\\MarketPilot_AI_Technical_Report.docx"
    doc.save(output_path)
    print(f"Report successfully saved to {output_path}")

if __name__ == "__main__":
    create_report()
